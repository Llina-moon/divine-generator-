from flask import Flask, request, send_file, render_template_string
from docx import Document
from docx.shared import Pt, RGBColor
import os

app = Flask(__name__)

def replace_text(doc, replacements):
    for p in doc.paragraphs:
        for run in p.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 0, 0)
    return doc

@app.route("/", methods=["GET"])
def form():
    return render_template_string(open("form.html", encoding="utf-8").read())

@app.route("/generate", methods=["POST"])
def generate():
    data = {
        "{ФИО}": request.form["fio"],
        "{Документ}": request.form["passport"],
        "{Адрес}": request.form["address"],
        "{Электронная почта лицензиара}": request.form["email"],
        "{Процент вознаграждения}": request.form["percent"],
        "{Процент вознаграждения прописью}": request.form["percent_words"],
        "{Срок действия договора}": request.form["duration"],
        "{Номер договора}": request.form["contract_id"],
        "{id artist / id contract}": request.form["contract_id"],
        "{Дата}": request.form["date"],
        "{ИНН}": request.form["inn"],
        "{СНИЛС}": request.form["snils"],
        "{Банковские реквизиты}": request.form["bank"]
    }

    fio = request.form["fio"].replace(" ", "_")

    doc1 = Document("contract_template.docx")
    doc1 = replace_text(doc1, data)
    contract_path = f"{fio}_договор.docx"
    doc1.save(contract_path)

    doc2 = Document("app_template.docx")
    doc2 = replace_text(doc2, data)
    app_path = f"{fio}_приложение.docx"
    doc2.save(app_path)

    return render_template_string(
        f"<h2>Файлы готовы!</h2>"
        f"<a href='/download/{contract_path}'>Скачать договор</a><br>"
        f"<a href='/download/{app_path}'>Скачать приложение</a>"
    )

@app.route("/download/<filename>")
def download(filename):
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)